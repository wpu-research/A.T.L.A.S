#!/usr/bin/env python3
"""
ATLAS DOCX Generator
Converts ATLAS_TR.md and ATLAS_EN.md to professional Word documents.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ---------------------------------------------------------------------------
# Color constants
# ---------------------------------------------------------------------------
COLOR_TABLE_HEADER_BG = RGBColor(0xBD, 0xD7, 0xEE)   # light blue
COLOR_CODE_BG         = RGBColor(0xF2, 0xF2, 0xF2)   # light gray
COLOR_HEADING1        = RGBColor(0x1F, 0x35, 0x64)   # dark navy
COLOR_HEADING2        = RGBColor(0x2E, 0x74, 0xB5)   # medium blue
COLOR_HEADING3        = RGBColor(0x40, 0x40, 0x40)   # dark gray


# ---------------------------------------------------------------------------
# Helpers – low-level XML
# ---------------------------------------------------------------------------

def rgb_to_hex(color: RGBColor) -> str:
    """Convert RGBColor (a tuple subclass) to a 6-digit hex string."""
    r, g, b = color[0], color[1], color[2]
    return f"{r:02X}{g:02X}{b:02X}"


def set_cell_shading(cell, fill_color: RGBColor):
    """Apply solid background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_to_hex(fill_color))
    tcPr.append(shd)


def set_paragraph_shading(paragraph, fill_color: RGBColor):
    """Apply solid background shading to a paragraph (for code blocks)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_to_hex(fill_color))
    pPr.append(shd)


def add_page_number_footer(document):
    """Insert a centered page-number field in the document's default footer."""
    section = document.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear any existing paragraphs
    for p in footer.paragraphs:
        p.clear()

    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    run.font.size = Pt(9)
    run.font.name = "Calibri"

    # PAGE field
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def set_page_margins(document, margin_cm: float = 2.5):
    """Set uniform page margins on all sections."""
    for section in document.sections:
        section.top_margin    = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin   = Cm(margin_cm)
        section.right_margin  = Cm(margin_cm)


# ---------------------------------------------------------------------------
# Inline-markup renderer
# ---------------------------------------------------------------------------

def add_inline_runs(paragraph, text: str, base_size: Pt = Pt(11),
                    base_bold: bool = False, base_font: str = "Calibri"):
    """
    Parse inline markdown (bold, inline-code) and add styled runs to paragraph.
    Handles: **bold**, `code`, and plain text.
    """
    # Pattern: **bold** or `code`
    pattern = re.compile(r"(\*\*(.+?)\*\*|`([^`]+?)`)")

    last = 0
    for m in pattern.finditer(text):
        # Plain text before match
        if m.start() > last:
            run = paragraph.add_run(text[last:m.start()])
            run.font.name = base_font
            run.font.size = base_size
            run.bold = base_bold

        if m.group(2) is not None:
            # Bold
            run = paragraph.add_run(m.group(2))
            run.font.name = base_font
            run.font.size = base_size
            run.bold = True
        elif m.group(3) is not None:
            # Inline code
            run = paragraph.add_run(m.group(3))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.bold = False

        last = m.end()

    # Remaining plain text
    if last < len(text):
        run = paragraph.add_run(text[last:])
        run.font.name = base_font
        run.font.size = base_size
        run.bold = base_bold


# ---------------------------------------------------------------------------
# Block-level elements
# ---------------------------------------------------------------------------

def add_title(document, text: str):
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(12)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = COLOR_HEADING1


def add_heading1(document, text: str):
    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = COLOR_HEADING1


def add_heading2(document, text: str):
    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = COLOR_HEADING2


def add_heading3(document, text: str):
    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(2)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = COLOR_HEADING3


def add_blockquote(document, text: str):
    """Render blockquote (> text) as indented italic paragraph."""
    para = document.add_paragraph()
    para.paragraph_format.left_indent  = Cm(1.0)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def add_body_paragraph(document, text: str):
    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(4)
    add_inline_runs(para, text)


def add_bullet(document, text: str, level: int = 0):
    para = document.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent   = Cm(0.5 + level * 0.5)
    para.paragraph_format.space_before  = Pt(1)
    para.paragraph_format.space_after   = Pt(1)
    add_inline_runs(para, text)


def add_code_block(document, lines: list):
    """Render a fenced code block with monospace font and gray background."""
    for line in lines:
        para = document.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.left_indent  = Cm(0.5)
        set_paragraph_shading(para, COLOR_CODE_BG)
        run = para.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    # Add a small spacer after the block
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after  = Pt(6)


def add_horizontal_rule(document):
    para = document.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table(document, headers: list, rows: list):
    """Render a markdown table as a Word table."""
    col_count = len(headers)
    table = document.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_shading(cell, COLOR_TABLE_HEADER_BG)
        para = cell.paragraphs[0]
        run = para.add_run(h.strip())
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.bold = True

    # Data rows
    for r_idx, row_data in enumerate(rows):
        tbl_row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= col_count:
                break
            cell = tbl_row.cells[c_idx]
            para = cell.paragraphs[0]
            add_inline_runs(para, cell_text.strip(), base_size=Pt(10))

    # Small spacer after table
    document.add_paragraph().paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# Markdown parser → Word
# ---------------------------------------------------------------------------

def parse_table_row(line: str) -> list:
    """Split a markdown table row into cells."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c for c in line.split("|")]


def is_separator_row(line: str) -> bool:
    """Detect table separator rows like |---|---|"""
    return bool(re.match(r"^\s*\|?\s*[-:]+\s*(\|\s*[-:]+\s*)+\|?\s*$", line))


def convert_markdown_to_docx(md_path: str, docx_path: str):
    """Parse markdown file and produce a Word document."""
    text = Path(md_path).read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    set_page_margins(doc, 2.5)
    add_page_number_footer(doc)

    # State
    in_code_block = False
    code_lines: list = []
    code_fence_pattern = re.compile(r"^```")

    # Table accumulator
    table_headers: list | None = None
    table_rows: list = []
    in_table = False

    i = 0
    while i < len(lines):
        line = lines[i]

        # ----------------------------------------------------------------
        # Code block
        # ----------------------------------------------------------------
        if code_fence_pattern.match(line):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                add_code_block(doc, code_lines)
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ----------------------------------------------------------------
        # Flush pending table if next line is not a table row
        # ----------------------------------------------------------------
        stripped = line.strip()
        is_table_line = stripped.startswith("|") or is_separator_row(stripped)

        if in_table and not is_table_line:
            if table_headers is not None and table_rows:
                add_table(doc, table_headers, table_rows)
            table_headers = None
            table_rows = []
            in_table = False

        # ----------------------------------------------------------------
        # Table rows
        # ----------------------------------------------------------------
        if is_table_line:
            if is_separator_row(stripped):
                # separator — skip
                i += 1
                continue
            cells = parse_table_row(stripped)
            if not in_table:
                # First row = headers
                table_headers = cells
                in_table = True
            else:
                table_rows.append(cells)
            i += 1
            continue

        # ----------------------------------------------------------------
        # Horizontal rule
        # ----------------------------------------------------------------
        if re.match(r"^---+\s*$", stripped) or re.match(r"^===+\s*$", stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ----------------------------------------------------------------
        # Headings
        # ----------------------------------------------------------------
        if stripped.startswith("#### "):
            add_heading3(doc, stripped[5:])
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading3(doc, stripped[4:])
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading2(doc, stripped[3:])
            i += 1
            continue
        if stripped.startswith("# "):
            # First # heading is the document title
            add_title(doc, stripped[2:])
            i += 1
            continue

        # ----------------------------------------------------------------
        # Blockquote
        # ----------------------------------------------------------------
        if stripped.startswith("> "):
            add_blockquote(doc, stripped[2:])
            i += 1
            continue

        # ----------------------------------------------------------------
        # Bullet list
        # ----------------------------------------------------------------
        bullet_m = re.match(r"^(\s*)[*\-] (.+)$", line)
        if bullet_m:
            indent_len = len(bullet_m.group(1))
            level = indent_len // 2
            add_bullet(doc, bullet_m.group(2), level)
            i += 1
            continue

        # Numbered list
        num_m = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if num_m:
            add_bullet(doc, num_m.group(1))
            i += 1
            continue

        # ----------------------------------------------------------------
        # Empty line
        # ----------------------------------------------------------------
        if stripped == "":
            i += 1
            continue

        # ----------------------------------------------------------------
        # Body paragraph
        # ----------------------------------------------------------------
        add_body_paragraph(doc, stripped)
        i += 1

    # Flush any remaining table
    if in_table and table_headers and table_rows:
        add_table(doc, table_headers, table_rows)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    base = Path("/home/wpu/Downloads/Projects/Atlas")

    convert_markdown_to_docx(
        str(base / "ATLAS_TR.md"),
        str(base / "ATLAS_TR.docx"),
    )
    convert_markdown_to_docx(
        str(base / "ATLAS_EN.md"),
        str(base / "ATLAS_EN.docx"),
    )
    print("Done. Both DOCX files generated successfully.")
