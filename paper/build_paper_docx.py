#!/usr/bin/env python3
"""Convert ATLAS_IEEE_Access_Draft.md to a formatted Word document."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(__file__).parent / "ATLAS_IEEE_Access_Draft.md"
OUT = Path(__file__).parent / "ATLAS_IEEE_Access_Draft.docx"

NAVY = RGBColor(0x1F, 0x35, 0x64)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
GRAY = RGBColor(0x40, 0x40, 0x40)
WARN = RGBColor(0xC0, 0x39, 0x1B)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def add_runs(par, text, base_size=10.5):
    for tok in INLINE.split(text):
        if not tok:
            continue
        run = par.add_run()
        run.font.size = Pt(base_size)
        if tok.startswith("**") and tok.endswith("**"):
            run.text = tok[2:-2]
            run.bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            run.text = tok[1:-1]
            run.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            run.text = tok[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        else:
            run.text = tok
        if "[MOCK" in tok or "[PLACEHOLDER" in tok or "MOCK DATA" in tok:
            run.font.color.rgb = WARN
            run.bold = True


def main():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.2)
        sec.left_margin = sec.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    lines = SRC.read_text(encoding="utf-8").split("\n")
    i = 0
    first_h1 = True
    while i < len(lines):
        line = lines[i]

        # tables
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            rows = []
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(header))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                p = cell.paragraphs[0]
                add_runs(p, h, 9.5)
                for r in p.runs:
                    r.bold = True
                shade_cell(cell, "BDD7EE")
            for ri, row in enumerate(rows):
                for j, val in enumerate(row[: len(header)]):
                    p = table.rows[ri + 1].cells[j].paragraphs[0]
                    add_runs(p, val, 9.5)
            doc.add_paragraph()
            continue

        if line.startswith("# ") :
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].replace("**", ""))
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = NAVY
            first_h1 = False
        elif line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].replace("**", ""))
            run.bold = True
            run.font.size = Pt(12.5)
            run.font.color.rgb = NAVY
            p.paragraph_format.space_before = Pt(14)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].replace("**", ""))
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = BLUE
            p.paragraph_format.space_before = Pt(10)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            add_runs(p, line[2:])
            for r in p.runs:
                r.font.color.rgb = WARN
        elif line.strip() == "---":
            pass
        elif re.match(r"^  \S", line):  # indented equation line
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, line.strip())
            for r in p.runs:
                if not r.text.strip().startswith("("):
                    r.italic = True
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            add_runs(p, line)
        elif line.strip():
            p = doc.add_paragraph()
            if not (line.startswith("[") and "]" in line[:6]):
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(p, line)
        i += 1

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
